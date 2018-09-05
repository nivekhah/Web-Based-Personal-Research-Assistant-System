# -*- coding:utf-8 -*-
from flask import Flask, render_template, request, jsonify, redirect, url_for, send_from_directory
from flask_login import LoginManager, login_required, login_user, current_user, logout_user
from werkzeug.security import generate_password_hash, check_password_hash
# from models import User
import csv
from docx import Document
import MySQLdb
import os
import base64
from DB import DB
from BibtexParser import BibtexParser
from EngName import EnglishName
from loadBibtex import LoadBibtex
import codecs
import json
import urllib
import re
import time
from flask_login import UserMixin
db = MySQLdb.connect(host='localhost', user='root', passwd='123456', db='test', port=3306, use_unicode=True,
                     charset="utf8")


class User(UserMixin):
    def __init__(self, id, username, password):
        self.id = id
        self.username = username
        self.password = password

    def __repr__(self):
        return '<User %r>' % self.username


login_manager = LoginManager()
app = Flask(__name__)
app.config['SECRET_KEY'] = '123456'
login_manager.session_protection='strong'  
login_manager.login_view = 'login'
login_manager.init_app(app)


@login_manager.user_loader
def load_user(user_id):
    cursor = db.cursor()
    # print(user_id)
    sql = 'select * from user where id = %s' % (user_id)
    cursor.execute(sql)
    datas = cursor.fetchall()
    user = User(datas[0][0], datas[0][1], datas[0][2])
    return user


# 用户注册页面的视图函数
@app.route('/register')
def register():
    return render_template('register.html')

# 接收用户注册信息的url视图
@app.route('/registerinfo')
def registerinfo():
    email = request.args.get('email')
    name = request.args.get('name')
    # idcard = request.args.get('idcard')
    phonenumber = request.args.get('phonenumber')
    password = request.args.get('password')
    # unit = request.args.get('unit')
    # 原始密码作为输入,以字符串形式输出密码的散列值,输出的值可保存在用户数据库中
    password_hash = generate_password_hash(password)
    cursor = db.cursor()
    sql = 'select email from user'
    cursor.execute(sql)
    datas = cursor.fetchall()
    list1 = list(datas)
    for l in list1:
        if l[0] == email:
            return '1'

    sql = 'select id from user'
    cursor.execute(sql)
    datas = cursor.fetchall()
    if len(datas)==0:
        sql = 'insert into user (id,email,name,phonenumber,password) values(%d,\'%s\',\'%s\',\'%s\',\'%s\')' % (
            1, email, name, phonenumber, password_hash)
    else:
        Max = max(datas)
        sql = 'insert into user (id,email,name,phonenumber,password) values(%d,\'%s\',\'%s\',\'%s\',\'%s\')' % (
            Max[0]+1, email, name, phonenumber, password_hash)
    # print check_password_hash(password_hash, email)
    cursor.execute(sql)
    db.commit()
    return '0'


# 对头像图片上传进行响应
@app.route('/upload/', methods=['POST', 'OPTIONS'])
def upload():
    if request.form.get("action") == "add":
        data = request.form.get("picStr")
        imgdata = base64.b64decode(data)
        # 使用用户名作为文件名字,每个用户的信息都存放在对应的文件夹下
        IMAGE_FOLDER = 'static/pic/' + current_user.username + '/'
        isExists = os.path.exists(IMAGE_FOLDER)
        if not isExists:
            os.makedirs(IMAGE_FOLDER)
        UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), IMAGE_FOLDER)
        imgfile=os.path.join(UPLOAD_FOLDER, "test.png")
        file=open(imgfile, 'wb')
        file.write(imgdata)
        file.close()
        db = DB('localhost', 'root', '123456', 'test', 3306)
        # 按照用户名进行插入
        db.InsertIntoPeopleByPicPath(IMAGE_FOLDER+'test.png',current_user.username)
        return jsonify(imgfile=imgfile)


@app.route('/')
def index():
    db = DB('localhost', 'root', '123456', 'test', 3306)
    number = 0
    if current_user.is_authenticated:
        print('已登录')
        datas = db.selectUsernameFromPeople()
        list1 = list(datas)
        for l in list1:
            # print(l[0])
            # print(current_user.username)
            if l[0] == current_user.username:
                number = 1
    else:
        print('未登录')
    return render_template('index.html', is_authenticated=current_user.is_authenticated, number=number)





@app.route('/pageone')
@login_required
def pageone():
    return render_template('pageone.html')


@app.route('/pagetwo')
@login_required
def pagetwo():
    return render_template('pagetwo.html')


@app.route('/pagethree')
@login_required
def pagethree():
    # 向添加头像界面传递用户名,以便更改的头像可以在改界面显示
    return render_template('pagethree.html', username=current_user.username)


@app.route('/pagefour', methods=['GET', 'POST'])
@login_required
def pagefour():
    if request.method == 'GET':
        return render_template('pagefour.html')
    else:
        db = DB('localhost', 'root', '123456', 'test', 3306)
        BibtexfilePath = './bib/' + current_user.username + '/'
        isExists = os.path.exists(BibtexfilePath)
        if not isExists:
            os.makedirs(BibtexfilePath)
        Bibtexfilename='total.bib'
        file = request.files['file']
        if file:
            file.save(BibtexfilePath + Bibtexfilename)
            bp = BibtexParser(BibtexfilePath + Bibtexfilename)
            try:
                keys = bp.getkey()
                for key in keys:
                    db.InsertIntoBibtexByUsername(current_user.username, key)
                return jsonify('0')
            except Exception:
                return jsonify('1')


@app.route('/pagefive')
@login_required
def pagefive():
    db = DB('localhost', 'root', '123456', 'test', 3306)
    results = db.selectAllFromPeopleByusername(current_user.username)
    domainName = results[4]
    return render_template('pagefive.html', domainName=domainName)


@app.route('/login', methods=['GET', 'POST'])
def login():
    email = request.args.get('email')
    password = request.args.get('password')
    if email == None and password == None:
        return render_template('login.html')
    else:
        cursor = db.cursor()
        sql = 'select * from user where email=\'%s\'' % (email)
        cursor.execute(sql)
        datas = cursor.fetchall()

        for data in datas:
            #print(data[1])
            user = User(data[0], data[1], data[4])
            if data[1] == email:
                if check_password_hash(data[4], password):
                    login_user(user)
                    return jsonify('Login successful')
                else:
                    return jsonify('Password error')
        # else:
        # return jsonify('No this user')
        return jsonify('No this user')


@app.route('/pageoneInfo')
def pageoneInfo():
    realname = request.args.get('realname')
    EnglishRealname = request.args.get('EnglishRealname')
    email = request.args.get('email')
    domainName = request.args.get('domainName')
    researchInterest = request.args.get('researchInterest')
    myself = request.args.get('myself')
    db = DB('localhost', 'root', '123456', 'test', 3306)
    datas = db.selectDomainFromPeople()
    list1 = list(datas)
    for l in list1:
        if l[0] == domainName:
            return jsonify('1')
    db.InsertIntoPeopleByBassicInfo(realname, EnglishRealname, email, current_user.username, domainName,
                                    researchInterest, myself)
    return jsonify('0')


@app.route('/pagetwoInfo')
def pagetwoInfo():
    Research_institutions = request.args.get('Research_institutions')
    department = request.args.get('department')
    position = request.args.get('position')
    # print(Research_institutions,department,position)
    db = DB('localhost', 'root', '123456', 'test', 3306)
    db.InsertIntoPeopleByWorkInfo(Research_institutions, department, position, current_user.username)

    return jsonify('0')

@app.route('/pagefourInfo')
def pagefourInfo():
    text = request.args.get('text')
    BibtexfilePath = './bib/' + current_user.username + '/'
    isExists = os.path.exists(BibtexfilePath)
    if not isExists:
        os.makedirs(BibtexfilePath)
    Bibtexfilename = 'total.bib'
    f = codecs.open(BibtexfilePath + Bibtexfilename, 'w', 'utf-8')
    f.write(text)
    f.close()
    bp = BibtexParser(BibtexfilePath + Bibtexfilename)
    keys = bp.getkey()
    for key in keys:
        db = DB('localhost', 'root', '123456', 'test', 3306)
        db.InsertIntoBibtexByUsername(current_user.username, key)
    return jsonify('0')


@app.route('/people/<domainName>')
def peopleshow(domainName):
    db = DB('localhost', 'root', '123456', 'test', 3306)
    results = db.selectAllFromPeople(domainName)
    realname = results[0]
    EnglishRealname = results[1]
    username=results[3]
    email = results[2]
    researchInterest = results[5]
    myself = results[6]
    Research_institutions = results[7]
    department = results[8]
    position = results[9]
    ProfilePicturePath = results[10]
    print(username)
    # 这里不能使用current_user.username
    data = db.SelectAllBibtexByUsername(username)
    BibtexfilePath = './bib/' + username + '/total.bib'
    if len(data) == 0:
        text = []
        path = []
        name = []
        keys = []
        return render_template('people.html', realname=realname, EnglishRealname=EnglishRealname,
                               ProfilePicturePath=ProfilePicturePath, Research_institutions=Research_institutions,
                               department=department, position=position, researchInterest=researchInterest,
                               myself=myself, email=email, name=name, text=text, path=path, keys=keys,
                               username=username)
    else:
        path = []
        bp = BibtexParser(BibtexfilePath)
        text, keys, titles, years = bp.gettextbykey()
        for key in keys:
            temp = db.SelectAllBibtexByUsernameAndbibtex(username, key)
            path.append(temp[0][2])
        return render_template('people.html', realname=realname, EnglishRealname=EnglishRealname,
                               ProfilePicturePath=ProfilePicturePath, Research_institutions=Research_institutions,
                               department=department, position=position, researchInterest=researchInterest,
                               myself=myself, email=email, text=text, path=path, name=titles, keys=keys,
                               username=username)


@app.route("/download/<username>/<filename>", methods=['GET'])
def download_file(username, filename):
    # 需要知道2个参数, 第1个参数是本地目录的path, 第2个参数是文件名(带扩展名)
    directory = os.getcwd()  # 假设在当前目录
    db = DB('localhost', 'root', '123456', 'test', 3306)
    # 这里不应该用current.username,应该用域名对应的username
    results = db.SelectAllBibtexByUsername(username)
    print(results)
    for result in results:
        temp = result[1]
        temp = temp.replace(":", "_")
        temp = temp.replace("/", "_")
        if (temp == filename):
            path = db.SelectAllBibtexByUsernameAndbibtex(username, result[1])[0][2]
            directory = directory+'/'+path.split('/')[1]+'/'+path.split('/')[2]+'/'
            file = path.split('/')[3]
            return send_from_directory(directory, file, as_attachment=True)


def sortbyyears(text, keys, titles, years):
    num = len(text)
    mytuple = []
    for i in range(0, num):
        mytuple.append([text[i], keys[i], titles[i], int(years[i])])
    print(len(text), len(keys), len(titles), len(years))
    print(mytuple)
    newmytuple=sorted(mytuple, key=lambda s: s[3], reverse=True)
    newtext = []
    newkesys = []
    newtitles = []
    newyears = []
    for i in range(0, num):
        newtext.append(newmytuple[i][0])
        newkesys.append(newmytuple[i][1])
        newtitles.append(newmytuple[i][2])
        newyears.append(newmytuple[i][3])
    return newtext, newkesys, newtitles, newyears





@app.route('/editpeople')
@login_required
def editpeople():
    db = DB('localhost', 'root', '123456', 'test', 3306)
    results = db.selectAllFromPeopleByusername(current_user.username)
    realname = results[0]
    EnglishRealname = results[1]
    email = results[2]
    researchInterest = results[5]
    myself = results[6]
    Research_institutions = results[7]
    department = results[8]
    position = results[9]
    ProfilePicturePath = results[10]
    data = db.SelectAllBibtexByUsername(current_user.username)
    BibtexfilePath = './bib/' + current_user.username + '/total.bib'
    if len(data) == 0:
        text = []
        path = []
        name = []
        keys = []
        return render_template('editpeople.html', realname=realname, EnglishRealname=EnglishRealname,
                               ProfilePicturePath=ProfilePicturePath, Research_institutions=Research_institutions,
                               department=department, position=position, researchInterest=researchInterest,
                               myself=myself, email=email, name=name, text=text, path=path, keys=keys,
                               username=current_user.username)
    else:
        path = []
        bp = BibtexParser(BibtexfilePath)
        text, keys, titles, years = bp.gettextbykey()
        text, keys, titles, years = sortbyyears(text, keys, titles, years)
        for key in keys:
            temp = db.SelectAllBibtexByUsernameAndbibtex(current_user.username, key)
            path.append(temp[0][2])
        return render_template('editpeople.html', realname=realname, EnglishRealname=EnglishRealname,
                               ProfilePicturePath=ProfilePicturePath, Research_institutions=Research_institutions,
                               department=department, position=position, researchInterest=researchInterest,
                               myself=myself, email=email, text=text, path=path, name=titles, keys=keys,
                               username=current_user.username)


@app.route('/edit')
@login_required
def edit():
    db = DB('localhost', 'root', '123456', 'test', 3306)
    results = db.selectAllFromPeopleByusername(current_user.username)
    realname = results[0]
    EnglishRealname = results[1]
    email = results[2]
    domainName = results[4]
    researchInterest = results[5]
    myself = results[6]
    Research_institutions = results[7]
    department = results[8]
    position = results[9]
    ProfilePicturePath = results[10]
    # bibtex = results[11]
    # print(bibtex)
    data = db.SelectAllBibtexByUsername(current_user.username)
    BibtexfilePath = './bib/' + current_user.username + '/total.bib'
    if len(data) == 0:
        keylist = []
        text = []
        return render_template('edit.html', realname=realname, EnglishRealname=EnglishRealname,
                               ProfilePicturePath=ProfilePicturePath, Research_institutions=Research_institutions,
                               department=department, position=position, researchInterest=researchInterest,
                               myself=myself, domainName=domainName,
                               email=email, text=text, keylist=keylist)
    else:
        # path=[]
        # name=[]
        bp = BibtexParser(BibtexfilePath)
        text, keys, titles, years = bp.gettextbykey()
        return render_template('edit.html', realname=realname, EnglishRealname=EnglishRealname,
                               ProfilePicturePath=ProfilePicturePath, Research_institutions=Research_institutions,
                               department=department, position=position, researchInterest=researchInterest,
                               myself=myself, email=email, keylist=keys, text=text, domainName=domainName)


@app.route('/editbibtex')
@login_required
def editbibtex():
    db = DB('localhost', 'root', '123456', 'test', 3306)
    # results = db.selectAllFromPeopleByusername(current_user.username)
    # bibtex = results[11]
    # print(bibtex)
    data = db.SelectAllBibtexByUsername(current_user.username)
    BibtexfilePath = './bib/' + current_user.username + '/total.bib'
    if len(data) == 0:
        keylist = []
        text = []
        return render_template('editbibtex.html', text=text, keylist=keylist)
    else:
        # path=[]
        # name=[]
        bp = BibtexParser(BibtexfilePath)
        text, keys, titles, years = bp.gettextbykey()
        return render_template('editbibtex.html', keylist=keys, text=text)


@app.route('/addbibtex')
@login_required
def addbibtex():
    db = DB('localhost', 'root', '123456', 'test', 3306)
    text = request.args.get('text')
    confirm = request.args.get('confirm')
    print(text)
    print(type(confirm))
    print(confirm)
    Bibtexfile = './bib/' + current_user.username + '/total.bib'
    bp = BibtexParser(Bibtexfile)
    oldbibtexkey = bp.getkey()

    BibtexfileTemp = './bib/' + current_user.username + '/temp.bib'
    f = codecs.open(BibtexfileTemp, 'w', 'utf-8')
    f.write(text)
    f.close()
    bptemp = BibtexParser(BibtexfileTemp)

    if confirm == '0':
        temp = bptemp.getTextTitle()[0]
        string = re.sub("[\s+\.\!\/_,$%^*\-(+\"\')]+|[+——()?{}【】“”！，。？、~@#￥%……&*（）]+", "", temp)
        print(string)
        texts, keys, titles, years = bp.gettextbykey()
        print(titles)
        print(len(titles))
        repeat = []
        for i in range(0, len(titles)):
            print(re.sub("[\s+\.\!\/_,$%^*\-(+\"\')]+|[+——()?{}【】“”！，。？、~@#￥%……&*（）]+", "", titles[i]))
            if re.sub("[\s+\.\!\/_,$%^*\-(+\"\')]+|[+——()?{}【】“”！，。？、~@#￥%……&*（）]+", "", titles[i]) == string:
                repeat.append(texts[i])
                # print()
                print(keys[i])
        print(repeat)
        if len(repeat) > 0:
            return jsonify('1', repeat)
        else:
            f = codecs.open(Bibtexfile, 'a', 'utf-8')
            f.write(text)
            f.close()
            db.InsertIntoBibtexByUsernameWithoutpdfpath(current_user.username, bptemp.getkey()[0])
            return jsonify(bptemp.getkey()[0])
    else:
        newbibtexkey = bptemp.getkey()[0]
        print(newbibtexkey)
        print(oldbibtexkey)
        for i in range(0, len(oldbibtexkey)):
            if newbibtexkey == oldbibtexkey[i]:
                return jsonify('2')
        db.InsertIntoBibtexByUsernameWithoutpdfpath(current_user.username, newbibtexkey)
        f = codecs.open(Bibtexfile, 'a', 'utf-8')
        f.write(text)
        f.close()
        return jsonify(newbibtexkey)



@app.route('/add', methods=['GET', 'POST'])
@login_required
def add():
    if request.method == 'GET':
        return render_template('add.html')
    else:
        file = request.files['file']
        text = request.form.get('mytext')
        path = './pdf/' + current_user.username + '/'
        isExists = os.path.exists(path)
        if not isExists:
            os.makedirs(path)
        temp = text.replace(":", "_")
        temp = temp.replace("/", "_")
        if file:
            file.save(path + temp+'.pdf')
            db = DB('localhost', 'root', '123456', 'test', 3306)
            db.UpdateBibtexByUsernameAndBibtex(path + temp+'.pdf', current_user.username, text)
            return '0'


@app.route('/save')
def save():
    realname = request.args.get('realname')
    EnglishRealname = request.args.get('EnglishRealname')
    email = request.args.get('email')
    Research_institutions = request.args.get('Research_institutions')
    department = request.args.get('department')
    position = request.args.get('position')
    domainName = request.args.get('domainName')
    researchInterest = request.args.get('researchInterest')
    myself = request.args.get('myself')
    # mylist = request.args.get('mylist')
    # reallist = json.loads(mylist)
    db = DB('localhost', 'root', '123456', 'test', 3306)
    db.UpdateBasicInfoByUsername(realname,EnglishRealname,email,Research_institutions,department,position,domainName,researchInterest,myself,current_user.username)
    return jsonify('0')


@app.route('/savebibtex')
def savebibtex():
    mylist = request.args.get('mylist')
    reallist = json.loads(mylist)
    db = DB('localhost', 'root', '123456', 'test', 3306)
    # bibtex数据中删除部分已经删除的数据
    # 将total.bib中的数据条目删除
    BibtexfilePath = './bib/' + current_user.username + '/total.bib'
    # BibtexfilePathnew = './bib/' + current_user.username + '/totalnew.bib'
    with open(BibtexfilePath, "r") as f:
        lines = f.readlines()
    with open(BibtexfilePath, "w",) as f_w:
        for i in range(0, len(reallist)):
            flag = 0
            for line in lines:
                if reallist[i] in line and '@' in line:
                    flag = 1
                temp = line.strip()
                temp = temp.replace('\n', '')
                if flag == 1:
                    f_w.write(line)
                if flag == 1 and temp == '}':
                    break

    datas = db.SelectAllBibtexByUsername(current_user.username)
    if len(datas) == 0:
        return jsonify('1')
    else:
        currentlist = []
        for data in datas:
            currentlist.append(data[1])
        differences = list(set(currentlist).difference(set(reallist)))
        for difference in differences:
            db.DeleteBibtexByUsernameAndBibtex(current_user.username, difference)
        return jsonify('0')


@app.route('/uploadpdf/<key>', methods=['GET', 'POST'])
def uploadpdf(key):
    if request.method == 'GET':
        return render_template("uploadpdf.html", key=key)
    else:
        file = request.files['file']
        path = './pdf/' + current_user.username + '/'
        isExists = os.path.exists(path)
        if not isExists:
            os.makedirs(path)
        key = urllib.quote(key.encode('utf8'))
        if file:
            file.save(path + key + '.pdf')
            db = DB('localhost', 'root', '123456', 'test', 3306)
            results = db.SelectAllBibtexByUsername(current_user.username)
            for result in results:
                temp = result[1]
                temp = temp.replace(":", "_")
                temp = temp.replace("/", "_")
                if(temp == key):
                    db.UpdateBibtexByUsernameAndBibtex(path + key + '.pdf', current_user.username, result[1])
        db = DB('localhost', 'root', '123456', 'test', 3306)
        results = db.selectAllFromPeopleByusername(current_user.username)
        domainName = results[4]
        return redirect(url_for('editpeople', domainName=domainName))


@app.route('/search')
def search():
    keyvalue = request.args.get('keyvalue')
    db = DB('localhost', 'root', '123456', 'test', 3306)
    datas = db.SelectAllFromPeopleByKeyValue(keyvalue)
    datas = list(datas)
    #for data in datas:
    #    print(data)
    return jsonify(datas)


@app.route('/table')
@login_required
def table():
    db = DB('localhost', 'root', '123456', 'test', 3306)
    ChineseName = db.selectAllFromPeopleByusername(current_user.username)
    # print(ChineseName[0])
    a = EnglishName(ChineseName[0])
    a.getAllName()
    return render_template('table.html',username = current_user.username)


@app.route('/sendFormat')
def sendFormat():
    db = DB('localhost', 'root', '123456', 'test', 3306)
    data = db.selectAllFromPeopleByusername(current_user.username)
    ChineseName = data[0]
    BibtexfilePath = './bib/' + current_user.username + '/total.bib'
    a = EnglishName(ChineseName)
    Columns = request.args.get('Columns')
    Columns = json.loads(Columns)
    types_num = request.args.get('types')
    order = request.args.get('order')
    order = json.loads(order)
    for i in range(len(Columns)):
        Columns[i] = Columns[i].strip('\n')
    if types_num == '1':
        types = 'article'
        b = LoadBibtex(BibtexfilePath, Columns, a.getAllName(), types)
        row_all = b.getInfo()
    elif types_num == '2':
        types = 'inproceedings'
        b = LoadBibtex(BibtexfilePath, Columns, a.getAllName(), types)
        row_all1 = b.getInfo()
        # types = 'proceedings'
        # c = LoadBibtex(BibtexfilePath, Columns, a.getAllName(), types)
        # row_all2 = c.getInfo()
        # row_all = row_all1 + row_all2
        row_all = row_all1
    else:
        types = 'inproceedings'
        b = LoadBibtex(BibtexfilePath, Columns, a.getAllName(), types)
        row_all1 = b.getInfo()
        # types = 'proceedings'
        # c = LoadBibtex(BibtexfilePath, Columns, a.getAllName(), types)
        # row_all2 = c.getInfo()
        types = 'article'
        d = LoadBibtex(BibtexfilePath, Columns, a.getAllName(), types)
        row_all3 = d.getInfo()
        row_all = row_all1 + row_all3
    for i in range(len(order)):
    # sorted(row_all, key=operator.itemgetter(1))
        key = order[len(order)-i-1].lower()
        if key == 'year':
            row_all.sort(key=lambda k: (k.get(key, 0)), reverse=True)
        else:
            row_all.sort(key=lambda k: (k.get(key, 0)))

    FilePath = './file/' + current_user.username + '/'
    isExists = os.path.exists(FilePath)
    if not isExists:
        os.makedirs(FilePath)

    mytime = time.strftime('%Y_%m_%d_%H_%M_%S',time.localtime(time.time()))
    document = Document()
    document.add_heading('Publications', 0)
    table = document.add_table(rows=1, cols=len(Columns), style='Table Grid')
    hdr_cells = table.rows[0].cells
    for i in range(len(Columns)):
        hdr_cells[i].text = Columns[i]
    for row in row_all:
        row_cells = table.add_row().cells
        for i in range(len(Columns)):
            if isinstance(row[Columns[i].lower()], int):
                field_value = str(row[Columns[i].lower()])
            else:
                field_value = row[Columns[i].lower()]
            row_cells[i].text = field_value

    document.add_page_break()

    document.save(FilePath+'pubs'+mytime+'.docx')

    for i in range(len(Columns)):
        Columns[i] = Columns[i].lower()
    #row_all.sort(key=lambda k: (k.get('title', 0)))
    with open(FilePath+'publications'+mytime+'.csv', 'w') as publication_csv_file:
        csv_writer = csv.DictWriter(publication_csv_file,
                                    fieldnames=Columns,
                                    delimiter=',')

        for row in row_all:
            if 'year' in Columns:

                csv_writer.writerow(row)
            else:
                row.pop('year')
                # print(row)
                csv_writer.writerow(row)
    publication_csv_file.close()
    return jsonify(row_all, mytime)


@app.route("/download/file/<username>/<filename>", methods=['GET'])
def downloadfile(username, filename):
    # 需要知道2个参数, 第1个参数是本地目录的path, 第2个参数是文件名(带扩展名)
    directory = './file/' + username
    return send_from_directory(directory, filename, as_attachment=True)


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))


if __name__ == '__main__':
    # app.run(debug=True, host='192.168.43.93', port=7000)
    app.run(debug=True, port=7000)
    # app.run(debug=True, host='192.168.1.109', port=7000)
    #app.run(debug=True, host='192.168.1.107', port=7000)


    # 添加论文时查重,依据什么查重,标题 去空格标点,大写小写 ,显示给用户,让用户判断是否重复
