import bibtexparser
from docx import Document
from docx.shared import Inches
# -*- coding: UTF-8 -*-

with open('Zeng-Deze3.bib') as bibtex_file:
    bibtex_str = bibtex_file.read()

# bibtex_str.replace('{', '')
# bibtex_str.replace('}', '')

bib_database = bibtexparser.loads(bibtex_str)

document = Document()
document.add_heading('Publications', 0)

all_row = []
for bib_entry in bib_database.entries:
    # print bib_entry
    if bib_entry['ENTRYTYPE'] == 'article':
        # print bib_entry
        row = {}
        if bib_entry.has_key('title'):
            row['title'] = bib_entry['title']
        if bib_entry.has_key('journal'):
            row['journal'] = bib_entry['journal']
        if bib_entry.has_key('year'):
            year = bib_entry['year']
        else:
            year = ''
        if bib_entry.has_key('number'):
            number = bib_entry['number']
        else:
            number = ''

        if bib_entry.has_key('volume'):
            volume = bib_entry['volume']
        else:
            volume = ''

        if bib_entry.has_key('pages'):
            pages = bib_entry['pages']
        else:
            pages = ''

        row['yearvolnopages'] = '%s, %s(%s), %s' % (year, number, volume, pages,)

        # print bib_entry

        authors = bib_entry['author'].replace(' and\n', ',')


        author_list= authors.split(',')
        print author_list
        print author_list.index('Deze Zeng')+1
        row['author'] = author_list.index('Deze Zeng')+1

        all_row.append(row)

table = document.add_table(rows=1, cols=4, style='Table Grid')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Title'
hdr_cells[1].text = 'Journal'
hdr_cells[2].text = 'Year, Vol(No)'
hdr_cells[3].text = 'Order'

all_row = sorted(all_row, key=lambda x:x['author'])



for row in all_row:
    row_cells = table.add_row().cells
    row_cells[0].text = row['title']
    row_cells[1].text = row['journal']
    row_cells[2].text = row['yearvolnopages']
    row_cells[3].text = str(row['author'])

document.add_page_break()

document.save('pubs.docx')
